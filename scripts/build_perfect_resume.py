"""Build recruiter-grade, ATS-optimized base resume (target: overall 10/10, ATS >= 9.5).

Rules applied as a recruiter would screen:
- No 'production support' / weak tenure branding
- Impact-first summary and bullets with metrics
- Single-column, Calibri, consistent sizes, standard headings
- Plain-text friendly for ATS parsers
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "shared" / "base_resume.docx"
ALSO = [
    ROOT / "PANKAJA_DevopsEngineer_Resume.docx",
    Path(r"C:\Users\pankajak\Downloads\PANKAJA_Resume_2YOE_Recruiter_Optimized.docx"),
]

NAME = "PANKAJA JAGDISH KULKARNI"
TITLE = "DevSecOps Engineer | Cloud Engineer | Cloud Security"
TAGLINE = "Azure  |  AKS  |  Kubernetes  |  DevSecOps  |  CI/CD  |  AZ-104 Certified"
CONTACT = (
    "Pune, India  |  +91 7045149611  |  pankajajagdishkulkarni@gmail.com  |  "
    "linkedin.com/in/pankaja-jagdish-kulkarni  |  github.com/Pankajajagdish  |  "
    "pankajajagdish.github.io"
)

SUMMARY = (
    "DevSecOps and Cloud Engineer at Amdocs delivering Azure and AKS platform reliability, "
    "cloud security, and automation for a telecom mediation platform. Built 8 internal tools "
    "that reduced manual operations effort by 40–90% and cut MTTR by 40%. Hands-on with "
    "Azure Policy, Microsoft Defender for Cloud, Entra ID RBAC, Terraform, Prometheus, "
    "Grafana, Azure Monitor, and GitHub Actions. Microsoft Certified: Azure Administrator "
    "Associate (AZ-104)."
)

SKILLS = [
    (
        "Cloud & Platform",
        "Microsoft Azure, Azure Kubernetes Service (AKS), AWS, Cloud Engineering, "
        "Virtual Networks (VNets), Private Endpoints, Key Vault, Azure Storage, Azure SQL, "
        "Managed Identity, Infrastructure as Code (IaC), Terraform, Azure Resource Manager (ARM), "
        "Cost Optimization",
    ),
    (
        "DevSecOps & Security",
        "DevSecOps, Cloud Security, Site Reliability Engineering (SRE), CI/CD, Azure DevOps, "
        "GitHub Actions, GitLab CI/CD, Docker, Helm, Microsoft Defender for Cloud, Azure Policy, "
        "Security Compliance, Entra ID, RBAC, IAM, PAM, MFA, CIS Benchmarks, NIST, CVSS, "
        "Vulnerability Management, NSGs, AWS IAM, AWS GuardDuty, AWS CloudTrail, AWS Config, "
        "S3 Security",
    ),
    (
        "SRE, Scripting & Automation",
        "Kubernetes, Node Pools, HPA, Prometheus, Grafana, Azure Monitor, Log Analytics, "
        "Observability, MTTR Reduction, Linux, Shell Scripting, Python, SQL, PostgreSQL, Git, "
        "Azure Functions, GitHub Copilot, Cursor IDE, Generative AI, AI Agents, MCP",
    ),
]

BULLETS = [
    "Drove Azure security compliance using Azure Policy, Defender for Cloud, and Entra ID RBAC on production subscriptions; weekly reports accelerated audit readiness and cut vulnerability triage time by 50%.",
    "Operated AKS clusters and node pools — tuned workloads, validated node versions pre-upgrade, and automated pod recovery for CrashLoopBackOff/OOMKilled; reduced MTTR by 40% and improved platform efficiency by 25%.",
    "Hardened Azure networking (VNets, NSGs, Private Endpoints, AMPLS) and executed AWS security reviews (IAM, S3, CloudTrail, GuardDuty) to align cross-cloud controls; improved incident resolution by 30%.",
    "Built and maintained CI/CD pipelines in Azure DevOps, GitHub Actions, and GitLab; pipeline optimization contributed to ~15% lower Azure infrastructure cost.",
    "Deployed observability with Prometheus, Grafana, Azure Monitor, and Log Analytics — increased coverage by 35% with MS Teams alerts for on-call response.",
    "Developed Python automation for PAM SLA tracking, daily health reporting, and Jira–Confluence sync; eliminated up to 90% of manual reporting and access-review effort.",
]

TOOLS = [
    ("Vulnerability Comparison Tool", "NIST/CVSS scan comparison; remediation prioritization improved 50%."),
    ("AKS Node Version Compliance Tool", "Pre-upgrade validation; manual checks reduced 60%."),
    ("Daily Health Report Agent & Git Release Diff Analyzer", "Automated ops reporting and release comparison."),
    ("Azure Cloud Security Auditor", "Defender, Azure Policy, Storage, Key Vault, and NSG audit automation."),
    ("Azure Cost Governance Reporter", "FinOps spend by resource group, untagged resources, CSV export."),
    ("Atlassian MCP SSO Toolkit", "Confluence and JIRA SSO MCP servers for Cursor IDE."),
]


def set_run(run, *, size=10, bold=False, name="Calibri", color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(p, before=0, after=4, line=1.08):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2F5496")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2, line=1.0)
    run = p.add_run(text.upper())
    set_run(run, size=11, bold=True, name="Calibri", color=RGBColor(0x1F, 0x4E, 0x79))
    add_horizontal_line(p)
    return p


def add_body(doc, text, *, size=10, bold=False, before=0, after=3):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after, line=1.08)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, name="Calibri")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=2, line=1.05)
    # clear default run and set formatting
    if p.runs:
        p.runs[0].text = text
        set_run(p.runs[0], size=10, bold=False, name="Calibri")
    else:
        run = p.add_run(text)
        set_run(run, size=10, bold=False, name="Calibri")
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_p, before=0, after=0, line=1.0)
    set_run(name_p.add_run(NAME), size=16, bold=True, name="Calibri", color=RGBColor(0x1F, 0x4E, 0x79))

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_p, before=2, after=0, line=1.0)
    set_run(title_p.add_run(TITLE), size=11, bold=True, name="Calibri")

    # Tagline (NO "2 years production..." branding)
    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(tag_p, before=2, after=2, line=1.0)
    set_run(tag_p.add_run(TAGLINE), size=9.5, bold=False, name="Calibri", color=RGBColor(0x40, 0x40, 0x40))

    # Contact
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact_p, before=0, after=4, line=1.05)
    set_run(contact_p.add_run(CONTACT), size=9, bold=False, name="Calibri")

    # Summary
    add_heading(doc, "Professional Summary")
    add_body(doc, SUMMARY, size=10, after=2)

    # Skills
    add_heading(doc, "Technical Skills")
    for label, items in SKILLS:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=1, after=2, line=1.05)
        set_run(p.add_run(f"{label}: "), size=10, bold=True, name="Calibri")
        set_run(p.add_run(items), size=10, bold=False, name="Calibri")

    # Experience
    add_heading(doc, "Professional Experience")
    add_body(doc, "Amdocs  |  Software Developer", size=10.5, bold=True, after=0)
    add_body(
        doc,
        "Pune, India  |  August 2024 – Present  |  Cloud, DevSecOps & SRE",
        size=9.5,
        after=1,
    )
    add_body(
        doc,
        "OneMediation — Telecom data processing platform on Azure / AKS",
        size=9.5,
        bold=True,
        after=2,
    )
    for b in BULLETS:
        add_bullet(doc, b)

    # Tools
    add_heading(doc, "Tools and Automation Built")
    for title, desc in TOOLS:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.05)
        set_run(p.add_run(f"{title} — "), size=10, bold=True, name="Calibri")
        set_run(p.add_run(desc), size=10, bold=False, name="Calibri")
    add_body(doc, "GitHub: https://github.com/Pankajajagdish", size=9.5, bold=False, after=2)

    # Certs
    add_heading(doc, "Certifications")
    add_body(
        doc,
        "Microsoft Certified: Azure Administrator Associate (AZ-104) — Active",
        size=10,
        after=2,
    )

    # Education
    add_heading(doc, "Education")
    add_body(
        doc,
        "B.E., Electronics and Telecommunication Engineering | Sinhgad College of Engineering, Pune | CGPA: 8.18/10 | 2024",
        size=10,
        after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    for dest in ALSO:
        try:
            doc.save(str(dest))
        except Exception as e:
            print(f"Skip {dest}: {e}")
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
