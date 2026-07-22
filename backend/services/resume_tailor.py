"""Resume tailor: copy your real resume, ONLY update TECHNICAL SKILLS for the JD.

Rules (strict):
- Keep every existing word in the resume — never delete or rewrite summary/experience/etc.
- Only modify the TECHNICAL SKILLS section: reorder JD-matched skills first + append missing JD skills.
- Target ATS / overall score >= 9.5 / 10.
"""

from __future__ import annotations

import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from services.profile import load_profile

ROOT = Path(__file__).resolve().parent.parent.parent
OUTPUT_DIR = Path(__file__).parent.parent / "output"
BASE_RESUME_CANDIDATES = [
    Path(__file__).parent.parent / "shared" / "base_resume.docx",
    ROOT / "shared" / "base_resume.docx",
    ROOT / "PANKAJA_DevopsEngineer_Resume.docx",
]

SKILLS_STOP_HEADINGS = {
    "PROFESSIONAL EXPERIENCE",
    "EXPERIENCE",
    "WORK EXPERIENCE",
    "TOOLS AND AUTOMATION BUILT",
    "PROJECTS",
    "CERTIFICATIONS",
    "EDUCATION",
    "PROFESSIONAL SUMMARY",
}

CATEGORY_HINTS = {
    "cloud": 0, "azure": 0, "aws": 0, "gcp": 0, "aks": 0, "terraform": 0,
    "arm": 0, "iac": 0, "vnet": 0, "networking": 0,
    "devops": 1, "devsecops": 1, "security": 1, "iam": 1, "rbac": 1,
    "entra": 1, "docker": 1, "helm": 1, "ci/cd": 1, "cicd": 1,
    "policy": 1, "defender": 1, "pam": 1,
    "kubernetes": 2, "k8s": 2, "prometheus": 2, "grafana": 2, "sre": 2,
    "python": 2, "linux": 2, "monitor": 2, "observability": 2, "shell": 2,
}


def _base_resume_path() -> Path:
    for p in BASE_RESUME_CANDIDATES:
        if p.exists():
            return p
    raise FileNotFoundError(
        "Base resume not found. Place your DOCX at shared/base_resume.docx"
    )


def extract_jd_keywords(jd: str, limit: int = 25) -> list[str]:
    profile = load_profile()
    jd_lower = (jd or "").lower()
    hits: list[str] = []

    for s in profile.get("skills", []):
        if s.lower() in jd_lower and s not in hits:
            hits.append(s)

    extras = [
        ("kubernetes", "Kubernetes"), ("k8s", "Kubernetes"), ("aks", "AKS"),
        ("terraform", "Terraform"), ("azure", "Microsoft Azure"), ("aws", "AWS"),
        ("devops", "DevOps"), ("devsecops", "DevSecOps"), ("ci/cd", "CI/CD"),
        ("cicd", "CI/CD"), ("docker", "Docker"), ("helm", "Helm"), ("iam", "IAM"),
        ("rbac", "RBAC"), ("entra", "Entra ID"), ("prometheus", "Prometheus"),
        ("grafana", "Grafana"), ("sre", "SRE"), ("python", "Python"),
        ("linux", "Linux"), ("ansible", "Ansible"), ("jenkins", "Jenkins"),
        ("argocd", "ArgoCD"), ("istio", "Istio"), ("openid", "OIDC"),
        ("oauth", "OAuth"), ("vault", "Key Vault"), ("key vault", "Key Vault"),
        ("defender", "Defender for Cloud"), ("azure policy", "Azure Policy"),
        ("github actions", "GitHub Actions"), ("azure devops", "Azure DevOps"),
        ("gitlab", "GitLab CI/CD"), ("cloud security", "Cloud Security"),
        ("platform engineer", "Platform Engineering"),
    ]
    for needle, label in extras:
        if needle in jd_lower and label not in hits:
            hits.append(label)
    return hits[:limit]


def _doc_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def compute_ats_score(text: str, jd_keywords: list[str]) -> float:
    if not jd_keywords:
        return 10.0
    t = text.lower()
    matched = sum(1 for k in jd_keywords if k and k.lower() in t)
    return round((matched / len(jd_keywords)) * 10.0, 2)


def compute_overall_score(text: str, jd_keywords: list[str], base_text: str) -> float:
    """Recruiter-style overall score (0–10): ATS fit + integrity + structure signals."""
    ats = compute_ats_score(text, jd_keywords)
    base_tokens = set(re.findall(r"[a-zA-Z0-9+#/.-]{3,}", base_text.lower()))
    new_tokens = set(re.findall(r"[a-zA-Z0-9+#/.-]{3,}", text.lower()))
    if not base_tokens:
        integrity = 10.0
    else:
        preserved = len(base_tokens & new_tokens) / len(base_tokens)
        integrity = round(preserved * 10.0, 2)

    required = [
        "professional summary",
        "technical skills",
        "professional experience",
        "certifications",
        "education",
        "amdocs",
        "az-104",
    ]
    t = text.lower()
    structure = round(10.0 * sum(1 for r in required if r in t) / len(required), 2)

    penalty = 0.0
    if "production support" in t or "2 years production experience" in t:
        penalty = 1.5

    overall = round(0.50 * ats + 0.30 * integrity + 0.20 * structure - penalty, 2)
    return min(10.0, max(0.0, overall))


def _is_heading(text: str) -> bool:
    t = (text or "").strip().upper()
    if not t:
        return False
    if t in SKILLS_STOP_HEADINGS or t == "TECHNICAL SKILLS":
        return True
    if t.isupper() and 3 <= len(t) <= 40 and ":" not in t:
        return True
    return False


def _find_skills_paragraphs(doc: Document) -> list[int]:
    start = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().upper() == "TECHNICAL SKILLS":
            start = i + 1
            break
    if start is None:
        return []
    indices = []
    for i in range(start, len(doc.paragraphs)):
        t = (doc.paragraphs[i].text or "").strip()
        if not t:
            continue
        if _is_heading(t) and t.upper() != "TECHNICAL SKILLS":
            break
        indices.append(i)
    return indices


def _split_skill_items(line: str) -> tuple[str, list[str]]:
    if ":" in line:
        label, rest = line.split(":", 1)
        return label.strip() + ": ", [x.strip() for x in rest.split(",") if x.strip()]
    return "", [x.strip() for x in line.split(",") if x.strip()]


def _already_covered(skill: str, existing: list[str]) -> bool:
    s = skill.lower()
    for e in existing:
        el = e.lower()
        if s == el or s in el or el in s:
            return True
    return False


def _category_for(skill: str, n_lines: int) -> int:
    sl = skill.lower()
    for hint, idx in CATEGORY_HINTS.items():
        if hint in sl:
            return min(idx, n_lines - 1)
    return min(1, n_lines - 1)


def _reorder_and_append(items: list[str], jd_keywords: list[str]) -> list[str]:
    jd_lower = [k.lower() for k in jd_keywords]

    def rank(item: str) -> tuple[int, int]:
        il = item.lower()
        for i, k in enumerate(jd_lower):
            if k in il or il in k:
                return (0, i)
        return (1, 0)

    indexed = list(enumerate(items))
    indexed.sort(key=lambda pair: (rank(pair[1])[0], rank(pair[1])[1], pair[0]))
    return [it for _, it in indexed]


def update_skills_only(doc: Document, jd: str) -> list[str]:
    jd_keywords = extract_jd_keywords(jd, 25)
    idxs = _find_skills_paragraphs(doc)
    if not idxs:
        raise RuntimeError("TECHNICAL SKILLS section not found in base resume")

    lines: list[tuple[str, list[str]]] = []
    for i in idxs:
        prefix, items = _split_skill_items(doc.paragraphs[i].text or "")
        lines.append((prefix, items))

    all_items = [it for _, items in lines for it in items]
    appended: list[str] = []
    for kw in jd_keywords:
        if _already_covered(kw, all_items):
            continue
        cat = _category_for(kw, len(lines))
        lines[cat][1].append(kw)
        all_items.append(kw)
        appended.append(kw)

    for li, (prefix, items) in enumerate(lines):
        lines[li] = (prefix, _reorder_and_append(items, jd_keywords))

    for para_i, (prefix, items) in zip(idxs, lines):
        new_text = prefix + ", ".join(items)
        para = doc.paragraphs[para_i]
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(new_text)
    return appended


def _boost_to_threshold(doc: Document, jd: str, base_text: str, min_score: float = 9.5) -> dict:
    keywords = extract_jd_keywords(jd, 25)
    text = _doc_text(doc)
    ats = compute_ats_score(text, keywords)
    overall = compute_overall_score(text, keywords, base_text)
    if ats >= min_score and overall >= min_score:
        return {"ats_score": ats, "overall_score": overall, "appended_extra": []}

    idxs = _find_skills_paragraphs(doc)
    if not idxs:
        return {"ats_score": ats, "overall_score": overall, "appended_extra": []}

    target_i = idxs[-1]
    prefix, items = _split_skill_items(doc.paragraphs[target_i].text or "")
    extra: list[str] = []

    for kw in keywords:
        text = _doc_text(doc)
        ats = compute_ats_score(text, keywords)
        overall = compute_overall_score(text, keywords, base_text)
        if ats >= min_score and overall >= min_score:
            break
        if kw.lower() in text.lower() or _already_covered(kw, items):
            continue
        items.append(kw)
        extra.append(kw)
        new_text = prefix + ", ".join(items)
        para = doc.paragraphs[target_i]
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(new_text)

    text = _doc_text(doc)
    return {
        "ats_score": compute_ats_score(text, keywords),
        "overall_score": compute_overall_score(text, keywords, base_text),
        "appended_extra": extra,
    }


def role_slug(job_title: str) -> str:
    """Sanitize job role for filename — role only, no company."""
    role = (job_title or "Cloud_Engineer").strip()
    role = re.sub(r"\([^)]*\)", " ", role)
    role = re.sub(r"\[[^\]]*\]", " ", role)
    role = re.split(r"\s+[-–—|@]\s+", role)[0]
    role = re.sub(r"[^\w\s+-]", " ", role)
    role = re.sub(r"\s+", "_", role.strip())
    role = re.sub(r"_+", "_", role).strip("_")
    return (role or "Cloud_Engineer")[:60]


def resume_filename_for_role(job_title: str) -> str:
    """e.g. PANKAJA_DevOps_Engineer.pdf — role only, never company."""
    return f"PANKAJA_{role_slug(job_title)}.pdf"


def convert_docx_to_pdf(docx_path: Path) -> Path:
    """Convert DOCX → PDF via LibreOffice (GitHub Actions / Linux) or Word (Windows)."""
    import subprocess
    import shutil as sh

    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")
    out_dir = docx_path.parent

    soffice = sh.which("soffice") or sh.which("libreoffice")
    if soffice:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            timeout=120,
            capture_output=True,
        )
        if not pdf_path.exists():
            raise RuntimeError(f"LibreOffice ran but PDF missing: {pdf_path}")
        return pdf_path

    # Windows fallback: Microsoft Word via docx2pdf (optional)
    try:
        from docx2pdf import convert as docx2pdf_convert  # type: ignore

        docx2pdf_convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception as e:
        print(f"docx2pdf fallback failed: {e}")

    raise RuntimeError(
        "PDF conversion failed. Install LibreOffice (soffice) or Word+docx2pdf."
    )


def generate_tailored_docx(job: dict, output_path: Path | None = None) -> str:
    """Build skills-tailored DOCX then convert to PDF. Returns PDF filename."""
    jd = job.get("description", "") or ""
    title = job.get("title", "Cloud Engineer")

    base_path = _base_resume_path()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    pdf_name = resume_filename_for_role(title)
    if output_path is None:
        pdf_path = OUTPUT_DIR / pdf_name
    else:
        pdf_path = Path(output_path)
        if pdf_path.suffix.lower() != ".pdf":
            pdf_path = pdf_path.with_suffix(".pdf")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # Work on a temp DOCX next to the final PDF
    docx_path = pdf_path.with_suffix(".docx")
    shutil.copy2(base_path, docx_path)
    doc = Document(str(docx_path))
    base_text = _doc_text(Document(str(base_path)))

    update_skills_only(doc, jd)
    scores = _boost_to_threshold(doc, jd, base_text, min_score=9.5)
    doc.save(str(docx_path))

    try:
        convert_docx_to_pdf(docx_path)
        if docx_path.exists():
            docx_path.unlink()
    except Exception as e:
        # Keep DOCX if PDF conversion fails so user still gets a file
        print(f"PDF conversion error ({e}); keeping DOCX")
        pdf_path = docx_path

    print(
        f"Resume tailored (skills only): ATS {scores['ats_score']}/10, "
        f"overall {scores['overall_score']}/10 -> {pdf_path.name}"
    )
    job["_ats_score"] = scores["ats_score"]
    job["_overall_score"] = scores["overall_score"]
    job["_resume_filename"] = pdf_path.name
    return pdf_path.name


def tailor_summary(jd: str, job_title: str) -> str:
    kws = extract_jd_keywords(jd, 6)
    top = ", ".join(kws[:5]) if kws else "Azure, Kubernetes, DevOps"
    return (
        f"Base resume kept as-is. TECHNICAL SKILLS updated for: {job_title}. "
        f"Emphasized: {top}."
    )


def generate_cover_letter_snippet(job: dict) -> str:
    profile = load_profile()
    jd = job.get("description", "")
    keywords = extract_jd_keywords(jd, 6)
    title = job.get("title", "the open position")
    company = job.get("company", "your organization")
    top = ", ".join(keywords[:4]) if keywords else "Azure, Kubernetes, DevOps, IAM"

    return (
        f"Dear Hiring Manager,\n\n"
        f"I am writing to apply for the {title} role at {company}. "
        f"I am a DevSecOps and Cloud Engineer at Amdocs with hands-on ownership of Azure/AKS platforms, "
        f"cloud security controls, and automation — including {top}.\n\n"
        f"Relevant outcomes from my work include reducing MTTR by 40% through AKS recovery automation, "
        f"cutting vulnerability triage time by 50%, and eliminating up to 90% of manual reporting with Python tooling. "
        f"I am Microsoft Certified: Azure Administrator Associate (AZ-104).\n\n"
        f"I would welcome the opportunity to discuss how I can contribute to {company}.\n\n"
        f"Best regards,\n"
        f"{profile['name']}\n"
        f"{profile['phone']} | {profile['email']}\n"
        f"{profile['linkedin']}"
    )


def generate_application_package(job: dict, app_dir: Path) -> dict:
    app_dir.mkdir(parents=True, exist_ok=True)
    jd = job.get("description", "")
    title = job.get("title", "")

    # Filename = role only (no company), e.g. PANKAJA_DevOps_Engineer.pdf
    resume_name = resume_filename_for_role(title)
    resume_path = app_dir / resume_name
    cover_path = app_dir / "cover_letter.txt"
    meta_path = app_dir / "meta.json"

    # Remove stale resumes (docx/pdf) so only the current role PDF remains
    for old in list(app_dir.glob("*.docx")) + list(app_dir.glob("*.pdf")):
        if old.name != resume_name:
            try:
                old.unlink()
            except OSError:
                pass

    actual_name = generate_tailored_docx(job, output_path=resume_path)
    resume_name = actual_name
    cover = generate_cover_letter_snippet(job)
    cover_path.write_text(cover, encoding="utf-8")

    note = tailor_summary(jd, title)
    meta = {
        "job_id": job["id"],
        "title": title,
        "company": job.get("company", ""),
        "url": job.get("url", ""),
        "match_score": job.get("match_score", 0),
        "posted_at": job.get("posted_at", ""),
        "tailored_summary": note,
        "ats_score": job.get("_ats_score"),
        "overall_score": job.get("_overall_score"),
        "resume_file": resume_name,
        "tailor_mode": "skills_only_keep_base_resume",
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }
    meta_path.write_text(json.dumps(meta, indent=2), encoding="utf-8")

    base = f"applications/{job['id']}"
    return {
        "resume_url": f"{base}/{resume_name}",
        "resume_file": resume_name,
        "cover_letter_url": f"{base}/cover_letter.txt",
        "cover_letter": cover,
        "tailored_summary": note,
        "ats_score": job.get("_ats_score"),
        "overall_score": job.get("_overall_score"),
        "application_ready": True,
    }
